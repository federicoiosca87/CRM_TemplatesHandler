"""
Word Document Parser for CMS Template Generator

Extracts structured content from localization Word documents.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from config import SECTION_MARKERS, TEMPLATE_VARIANTS, LANGUAGE_MAPPING

_BBCODE_TAG_RE = re.compile(r'\[/?(?:b|i|u|url(?:=[^\]]*)?)\]')


def _strip_bbcode(text: str) -> str:
    """Remove BBCode formatting tags for plain-text matching."""
    return _BBCODE_TAG_RE.sub('', text)


def _bullets_to_bbcode_list(text: str) -> str:
    """Convert consecutive lines starting with '• ' into [ul][li]...[/li][/ul] BBCode."""
    if not text or '• ' not in text:
        return text
    lines = text.split('\n')
    result = []
    in_list = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('• '):
            if not in_list:
                result.append('[ul]')
                in_list = True
            content = stripped[2:]  # remove '• ' prefix
            result.append(f'[li]{content}[/li]')
        else:
            if in_list:
                result.append('[/ul]')
                in_list = False
            result.append(line)
    if in_list:
        result.append('[/ul]')
    return '\n'.join(result)


def _run_to_bbcode(run) -> str:
    """Convert a single Word run to BBCode-formatted text."""
    text = run.text
    if not text:
        return ""
    if run.underline:
        text = f"[u]{text}[/u]"
    if run.italic is True:
        text = f"[i]{text}[/i]"
    if run.bold is True:
        text = f"[b]{text}[/b]"
    return text


def _paragraph_to_bbcode(para) -> str:
    """Convert a Word paragraph's runs to text with BBCode formatting.

    Preserves explicit bold, italic, and underline formatting applied in Word
    as [b], [i], [u] BBCode tags.  Hyperlinks are converted to
    [url=...]text[/url] BBCode.  Style-inherited formatting (e.g. a heading
    that is bold by virtue of its style) is ignored so that section headers
    remain plain text for matching purposes.
    """
    p_element = para._p
    hyperlink_tag = qn('w:hyperlink')
    run_tag = qn('w:r')

    has_hyperlinks = p_element.findall(hyperlink_tag)

    # Fast path: no hyperlinks — use simple run iteration
    if not has_hyperlinks:
        if not para.runs:
            return para.text.strip()

        parts = []
        for run in para.runs:
            parts.append(_run_to_bbcode(run))

        result = "".join(parts).strip()
        for tag in ('b', 'i', 'u'):
            result = result.replace(f"[/{tag}][{tag}]", "")
        return result

    # Walk XML children to preserve hyperlink boundaries
    from docx.text.run import Run
    parts = []
    rels = para.part.rels

    for child in p_element:
        if child.tag == run_tag:
            run = Run(child, para)
            parts.append(_run_to_bbcode(run))
        elif child.tag == hyperlink_tag:
            # Extract URL from relationship
            r_id = child.get(qn('r:id'))
            url = ""
            if r_id and r_id in rels:
                url = rels[r_id].target_ref
            # Collect text from runs inside the hyperlink
            link_parts = []
            for sub in child.findall(run_tag):
                run = Run(sub, para)
                link_parts.append(_run_to_bbcode(run))
            link_text = "".join(link_parts)
            if url and link_text:
                parts.append(f"[url={url}]{link_text}[/url]")
            elif link_text:
                parts.append(link_text)

    result = "".join(parts).strip()
    for tag in ('b', 'i', 'u'):
        result = result.replace(f"[/{tag}][{tag}]", "")
    return result

# Localized field labels (Title / Body / CTA equivalents)
# Icelandic: Titill (Title), Meginmál (Body)
# Estonian: Pealkiri (Title)
# Spanish: Título/Titulo (Title), Cuerpo (Body)
TITLE_LABELS = {"TITLE", "PEALKIRI", "TITILL", "TÍTULO", "TITULO"}
BODY_LABELS = {"BODY", "MEGINMÁL", "MEGINMAL", "CUERPO"}
CTA_LABELS = {"CTA", "CALL TO ACTION", "CALLTOACTION"}
ACTION_KEY_LABELS = {"ACTION KEY", "ACTIONKEY", "ACTION"}
ACTION_VALUE_LABELS = {"ACTION VALUE", "ACTIONVALUE"}


def _is_label_start(text_upper: str, labels: set[str]) -> bool:
    """Check if text starts with a label followed by newline or colon+newline."""
    for label in labels:
        if text_upper.startswith(label + "\n") or text_upper.startswith(label + ":\n"):
            return True
    return False


def _is_standalone_label(text_upper: str, labels: set[str]) -> bool:
    """Check if text is exactly a label (with optional colon)."""
    return text_upper in labels or (text_upper.rstrip(":") in labels and text_upper.endswith(":"))


def _strip_label_prefix(text: str, text_upper: str, labels: set[str]) -> Optional[str]:
    """If text starts with 'Label: content' or 'Label content', return the content after the label."""
    for label in labels:
        if text_upper.startswith(label + ":") or text_upper.startswith(label + " "):
            rest = text[len(label):].lstrip(": ").strip()
            return rest if rest else None
    return None


@dataclass
class TemplateContent:
    """Content for a single template variant."""
    variant: str
    title: Optional[str] = None
    body: Optional[str] = None
    cta: Optional[str] = None
    cta_mobile: Optional[str] = None
    action_key: Optional[str] = None
    action_value: Optional[str] = None
    send_condition: str = "JoinedCampaign"  # JoinedCampaign (Launch) or NotOptedIn (Reminder)


@dataclass
class OmsSection:
    """OMS section (Launch or Reminder) with multiple template variants."""
    section_type: str  # "Launch" or "Reminder"
    templates: list[TemplateContent] = field(default_factory=list)


@dataclass
class SmsSection:
    """SMS section (Launch or Reminder) with multiple template variants."""
    section_type: str  # "Launch" or "Reminder"
    templates: list[TemplateContent] = field(default_factory=list)


@dataclass
class PushSection:
    """Push notification section (Launch, Reminder, or Reward) with multiple template variants."""
    section_type: str  # "Launch", "Reminder", or "Reward"
    templates: list[TemplateContent] = field(default_factory=list)


@dataclass
class TcSection:
    """Terms & Conditions section."""
    significant_terms: Optional[str] = None
    terms_and_conditions: Optional[str] = None


@dataclass
class MyOffersSection:
    """My Offers inbox content."""
    headline: Optional[str] = None
    sub_headline: Optional[str] = None
    task: Optional[str] = None
    reward: Optional[str] = None


@dataclass
class ParsedDocument:
    """Complete parsed content from a Word document."""
    language_code: str
    offer_name: str
    my_offers: Optional[MyOffersSection] = None
    launch_oms: Optional[OmsSection] = None
    reminder_oms: Optional[OmsSection] = None
    reward_oms: Optional[OmsSection] = None
    launch_sms: Optional[SmsSection] = None
    reminder_sms: Optional[SmsSection] = None
    launch_push: Optional[PushSection] = None
    reminder_push: Optional[PushSection] = None
    reward_push: Optional[PushSection] = None
    tc: Optional[TcSection] = None
    raw_paragraphs: list[str] = field(default_factory=list)


def extract_language_and_offer(filename: str) -> tuple[str, str]:
    """
    Extract language code and offer name from filename.
    Format: {LANGUAGE}_{offer_name}.docx
    
    Handles compound language codes like EN_PE, RU_ET by checking
    against known LANGUAGE_MAPPING keys.
    """
    stem = Path(filename).stem
    stem_upper = stem.upper()
    
    # Try to find longest matching language code from LANGUAGE_MAPPING
    # Sort by length descending to match longer codes first (e.g., EN_PE before EN)
    known_codes = sorted(LANGUAGE_MAPPING.keys(), key=len, reverse=True)
    
    for code in known_codes:
        if stem_upper.startswith(code + "_"):
            # Found a match - the rest is the offer name
            offer_name = stem[len(code) + 1:]  # +1 for underscore
            return code, offer_name
    
    # Fallback: split on first underscore
    parts = stem.split("_", 1)
    if len(parts) == 2:
        return parts[0].upper(), parts[1]
    return stem.upper(), stem


def parse_word_document(file_path: Path) -> ParsedDocument:
    """
    Parse a Word document and extract structured content.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        ParsedDocument with all extracted sections
    """
    doc = Document(file_path)
    language_code, offer_name = extract_language_and_offer(file_path.name)
    
    # Extract all paragraphs with their text, preserving bullet list formatting.
    # Also include table-cell text because many templates are authored in tables.
    paragraphs = []
    for para in doc.paragraphs:
        text = _paragraph_to_bbcode(para)
        if text:
            # Check if this is a list item (bullet or numbered)
            is_list_item = False
            if para._p.pPr is not None and para._p.pPr.numPr is not None:
                is_list_item = True
            else:
                try:
                    if para.style and 'list' in para.style.name.lower():
                        is_list_item = True
                except Exception:
                    pass  # Skip style check for corrupted docs
            
            # Prefix list items with bullet character
            if is_list_item:
                text = "• " + text
            
            paragraphs.append(text)

    # Capture text authored inside tables (cells are not included in doc.paragraphs).
    # Table-based docs use key-value rows: [Label, Value]. Merged cells produce
    # duplicate text across columns, so we deduplicate per-row.
    _CMS_METADATA_RE = re.compile(
        r"^CampaignWizard(Oms|Sms|TC)Template\.", re.IGNORECASE
    )
    for table in doc.tables:
        for row in table.rows:
            seen_in_row = set()
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = _paragraph_to_bbcode(para)
                    # Normalize non-breaking spaces
                    text = text.replace("\xa0", " ")
                    if not text:
                        continue
                    # Skip CMS metadata rows (e.g. "CampaignWizardOmsTemplate.Deposit-...")
                    plain_text = _strip_bbcode(text)
                    if _CMS_METADATA_RE.match(plain_text):
                        continue
                    # Deduplicate merged cells within the same row
                    if plain_text in seen_in_row:
                        continue
                    seen_in_row.add(plain_text)

                    is_list_item = False
                    if para._p.pPr is not None and para._p.pPr.numPr is not None:
                        is_list_item = True
                    else:
                        try:
                            if para.style and "list" in para.style.name.lower():
                                is_list_item = True
                        except Exception:
                            pass  # Skip style check for corrupted docs

                    if is_list_item:
                        text = "• " + text

                    paragraphs.append(text)
    
    parsed = ParsedDocument(
        language_code=language_code,
        offer_name=offer_name,
        raw_paragraphs=paragraphs,
    )
    
    # Parse sections based on markers
    parsed.my_offers = _parse_my_offers_section(paragraphs)
    parsed.launch_oms = _parse_oms_section(paragraphs, "Launch")
    parsed.reminder_oms = _parse_oms_section(paragraphs, "Reminder")
    parsed.reward_oms = _parse_reward_oms_section(paragraphs)
    parsed.launch_sms = _parse_sms_section(paragraphs, "Launch")
    parsed.reminder_sms = _parse_sms_section(paragraphs, "Reminder")
    parsed.launch_push = _parse_push_section(paragraphs, "Launch")
    parsed.reminder_push = _parse_push_section(paragraphs, "Reminder")
    parsed.reward_push = _parse_reward_push_section(paragraphs)
    parsed.tc = _parse_tc_section(paragraphs)
    
    return parsed


def _find_section_start(paragraphs: list[str], markers: list[str]) -> int:
    """Find the index where a section starts based on markers."""
    for i, para in enumerate(paragraphs):
        para_upper = _strip_bbcode(para).upper().strip()
        for marker in markers:
            if marker in para_upper:
                return i
    return -1


def _find_tc_section_start(paragraphs: list[str], markers: list[str]) -> int:
    """Find T&C section start using strict matching (exact line or line starts with marker).
    
    This is stricter than _find_section_start because TC markers like 'TINGIMUSED'
    might appear as substrings in body text (e.g., 'Kehtivad tingimused. 21+.').
    """
    for i, para in enumerate(paragraphs):
        para_upper = _strip_bbcode(para).upper().strip()
        for marker in markers:
            # Exact match (standalone header)
            if para_upper == marker:
                return i
            # Line starts with marker followed by nothing, space, or punctuation
            if para_upper.startswith(marker):
                rest = para_upper[len(marker):]
                if not rest or rest[0] in ' \t:–-':
                    return i
    return -1


def _parse_my_offers_section(paragraphs: list[str]) -> Optional[MyOffersSection]:
    """Parse MY OFFERS section for inbox content."""
    start = _find_section_start(paragraphs, SECTION_MARKERS["MY_OFFERS"])
    if start == -1:
        return None
    
    section = MyOffersSection()
    current_field = None
    
    # Exact T&C section headers for boundary detection
    tc_exact_headers = [
        "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
        "ΟΡΟΙ ΚΑΙ ΠΡΟΫΠΟΘΕΣΕΙΣ", "ΣΗΜΑΝΤΙΚΟΙ ΟΡΟΙ", "ΠΛΗΡΕΙΣ ΟΡΟΙ",
        "TÉRMINOS Y CONDICIONES", "TERMOS E CONDIÇÕES",
        "TERMINI E CONDIZIONI", "CONDITIONS GÉNÉRALES", "AGB",
        "TINGIMUSED", "OLULISED TINGIMUSED", "TÄIELIKUD TINGIMUSED",  # Estonian
        "REEGLID JA TINGIMUSED", "REEGLID", "OLULISED REEGLID",  # Estonian (alt)
        "REGLUR OG SKILYRÐI", "SKILYRÐI",  # Icelandic
        "KÄYTTÖEHDOT", "EHDOT",  # Finnish
        "TERMS & CONDITIONS",  # English variant
        "ŞARTLAR VE KOŞULLAR", "ŞARTLAR & KOŞULLAR",  # Turkish
        "ÖNEMLI ŞARTLAR VE KOŞULLAR",  # Turkish (significant)
        "WARUNKI I ZASADY", "ISTOTNE WARUNKI I ZASADY",  # Polish
        "BETINGELSER OG VILKÅR", "VIKTIGE BETINGELSER OG VILKÅR",  # Norwegian
        "NOTEIKUMI", "BŪTISKIE NOTEIKUMI",  # Latvian
    ]
    
    for para in paragraphs[start:]:
        para_upper = _strip_bbcode(para).upper().strip()
        
        # Check for next major section
        if "LAUNCH OMS" in para_upper or ("SMS" in para_upper and "OMS" not in para_upper):
            break
        if para_upper in tc_exact_headers:
            break
            
        if "HEADLINE" in para_upper and "SUB" not in para_upper:
            current_field = "headline"
        elif "SUB-HEADLINE" in para_upper or "SUBHEADLINE" in para_upper:
            current_field = "sub_headline"
        elif "TASK" in para_upper:
            current_field = "task"
        elif "REWARD" in para_upper:
            current_field = "reward"
        elif current_field and para.strip():
            # Append content to current field
            current_value = getattr(section, current_field) or ""
            setattr(section, current_field, (current_value + " " + para).strip())
    
    return section


def _parse_oms_section(paragraphs: list[str], section_type: str) -> Optional[OmsSection]:
    """Parse OMS section (Launch or Reminder) with template variants."""
    marker_key = f"{section_type.upper()}_OMS"
    markers = SECTION_MARKERS.get(marker_key, [section_type.upper()])
    
    start = _find_section_start(paragraphs, markers)
    if start == -1:
        return None
    
    # Map section type to SendCondition
    send_condition = "JoinedCampaign" if section_type == "Launch" else "NotOptedIn"
    
    section = OmsSection(section_type=section_type)
    current_template = None
    current_field = None
    
    # Find end of section
    end = len(paragraphs)
    for i, para in enumerate(paragraphs[start + 1:], start + 1):
        para_upper = _strip_bbcode(para).upper().strip()
        # Check if we hit another major section
        if section_type == "Launch" and "REMINDER" in para_upper:
            end = i
            break
        elif "REWARD RECEIVED" in para_upper:
            # End of templates, reward received section starts
            end = i
            break
        elif "SMS" in para_upper and "OMS" not in para_upper:
            end = i
            break
        # Use exact T&C section headers only
        tc_exact_headers = [
            "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
            "ΟΡΟΙ ΚΑΙ ΠΡΟΫΠΟΘΕΣΕΙΣ", "ΣΗΜΑΝΤΙΚΟΙ ΟΡΟΙ", "ΠΛΗΡΕΙΣ ΟΡΟΙ",
            "TÉRMINOS Y CONDICIONES", "TERMOS E CONDIÇÕES",
            "TERMINI E CONDIZIONI", "CONDITIONS GÉNÉRALES", "AGB",
            "TINGIMUSED", "OLULISED TINGIMUSED", "TÄIELIKUD TINGIMUSED",  # Estonian
            "REEGLID JA TINGIMUSED", "REEGLID", "OLULISED REEGLID",  # Estonian (alt)
            "REGLUR OG SKILYRÐI", "SKILYRÐI",  # Icelandic
            "KÄYTTÖEHDOT", "EHDOT",  # Finnish
            "TERMS & CONDITIONS",  # English variant
            "ŞARTLAR VE KOŞULLAR", "ŞARTLAR & KOŞULLAR",  # Turkish
            "ÖNEMLI ŞARTLAR VE KOŞULLAR",  # Turkish (significant)
            "WARUNKI I ZASADY", "ISTOTNE WARUNKI I ZASADY",  # Polish
            "BETINGELSER OG VILKÅR", "VIKTIGE BETINGELSER OG VILKÅR",  # Norwegian
            "NOTEIKUMI", "BŪTISKIE NOTEIKUMI",  # Latvian
        ]
        if para_upper in tc_exact_headers:
            end = i
            break
    
    for para in paragraphs[start:end]:
        para_stripped = para.strip()
        para_upper = _strip_bbcode(para_stripped).upper()
        
        # Check for combined section+variant header (table-based docs)
        # e.g. "LAUNCH OMS - TemplateA", "REMINDER OMS - TemplateB"
        # Also handles Estonian: "Näidis" (Template), "Meeldetuletus" (Reminder), "Lansseerimine" (Launch)
        combined_match = re.search(r"(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])", para_upper)
        if combined_match and ("LAUNCH" in para_upper or "REMINDER" in para_upper or "OMS" in para_upper or "MEELDETULETUS" in para_upper or "LANSSEERIMINE" in para_upper):
            if current_template:
                section.templates.append(current_template)
            current_template = TemplateContent(
                variant=combined_match.group(1),
                send_condition=send_condition
            )
            current_field = "title"
            continue

        # Check for "Launch A" / "Reminder A" format (DK-style docs)
        # e.g. "Launch A", "Reminder B" — section type + variant letter, no "Template" keyword
        # Also handles Estonian equivalents: "Lansseerimine A", "Meeldetuletus B"
        short_variant_match = re.match(
            r"^(?:LAUNCH|REMINDER|LANSSEERIMINE|MEELDETULETUS|OMS\s+(?:LAUNCH|REMINDER))\s+([A-F])$", para_upper
        )
        if short_variant_match:
            if current_template:
                section.templates.append(current_template)
            current_template = TemplateContent(
                variant=short_variant_match.group(1),
                send_condition=send_condition
            )
            current_field = "title"
            continue

        # Check for template variant marker
        variant_match = re.match(r"(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])", para_upper)
        if variant_match:
            if current_template:
                section.templates.append(current_template)
            current_template = TemplateContent(
                variant=variant_match.group(1),
                send_condition=send_condition
            )
            # Default to title field - content before BODY/CTA is the title
            # This handles docs with no explicit "Title" label
            current_field = "title"
            continue
        
        if current_template:
            # Check for field labels (explicit markers)
            # Handle various formats: "TITLE" alone, "TITLE:", "Title\n...", etc.
            
            # Check if paragraph starts with a label followed by newline (embedded label)
            # e.g., "Title\n🎮 Bet on Sports..." or "Body\n🎯 Get..."
            para_plain_upper = _strip_bbcode(para_stripped).upper()
            if _is_label_start(para_plain_upper, TITLE_LABELS):
                current_field = "title"
                # Extract content after the label
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.title = (current_template.title or "") + "\n" + rest if current_template.title else rest
                        current_template.title = current_template.title.strip()
                continue
            elif _is_label_start(para_plain_upper, BODY_LABELS):
                current_field = "body"
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.body = (current_template.body or "") + "\n" + rest if current_template.body else rest
                        current_template.body = current_template.body.strip()
                continue
            elif _is_label_start(para_plain_upper, CTA_LABELS):
                current_field = "cta"
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.cta = (current_template.cta or "") + "\n" + rest if current_template.cta else rest
                        current_template.cta = current_template.cta.strip()
                continue
            
            # Handle standalone labels (no embedded content)
            if _is_standalone_label(para_upper, TITLE_LABELS):
                current_field = "title"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS) is not None:
                current_field = "title"
                rest = _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS)
                if rest:
                    current_template.title = rest
                continue
            elif _is_standalone_label(para_upper, BODY_LABELS):
                current_field = "body"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, BODY_LABELS) is not None:
                current_field = "body"
                rest = _strip_label_prefix(para_stripped, para_upper, BODY_LABELS)
                if rest:
                    current_template.body = rest
                continue
            elif _is_standalone_label(para_upper, CTA_LABELS):
                current_field = "cta"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, CTA_LABELS) is not None:
                current_field = "cta"
                rest = _strip_label_prefix(para_stripped, para_upper, CTA_LABELS)
                if rest:
                    current_template.cta = rest
                continue
            elif current_field and para_stripped:
                current_value = getattr(current_template, current_field) or ""
                new_value = (current_value + "\n" + para_stripped).strip()
                setattr(current_template, current_field, new_value)
    
    # Don't forget the last template
    if current_template:
        section.templates.append(current_template)
    
    return section if section.templates else None


def _parse_reward_oms_section(paragraphs: list[str]) -> Optional[OmsSection]:
    """Parse REWARD RECEIVED OMS section with template variants."""
    markers = SECTION_MARKERS.get("REWARD_OMS", ["REWARD RECEIVED"])
    
    start = _find_section_start(paragraphs, markers)
    if start == -1:
        return None
    
    section = OmsSection(section_type="Reward")
    current_template = None
    current_field = None
    
    # Check if header itself contains template variant (e.g., "REWARD RECEIVED – OMS – Template A")
    header_para = _strip_bbcode(paragraphs[start]).upper()
    header_variant = None
    variant_match = re.search(r'(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])', header_para)
    if variant_match:
        header_variant = variant_match.group(1)
    
    # Find end of section - ends at T&C or SMS or next major section
    # Use exact T&C section headers only (not partial matches)
    tc_exact_headers = [
        "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
        "ΟΡΟΙ ΚΑΙ ΠΡΟΫΠΟΘΕΣΕΙΣ", "ΣΗΜΑΝΤΙΚΟΙ ΟΡΟΙ", "ΠΛΗΡΕΙΣ ΟΡΟΙ",
        "TÉRMINOS Y CONDICIONES", "TERMOS E CONDIÇÕES",
        "TERMINI E CONDIZIONI", "CONDITIONS GÉNÉRALES", "AGB",
        "TINGIMUSED", "OLULISED TINGIMUSED", "TÄIELIKUD TINGIMUSED",  # Estonian
        "REEGLID JA TINGIMUSED", "REEGLID", "OLULISED REEGLID",  # Estonian (alt)
        "REGLUR OG SKILYRÐI", "SKILYRÐI",  # Icelandic
        "KÄYTTÖEHDOT", "EHDOT",  # Finnish
        "TERMS & CONDITIONS",  # English variant
        "ŞARTLAR VE KOŞULLAR", "ŞARTLAR & KOŞULLAR",  # Turkish
        "ÖNEMLI ŞARTLAR VE KOŞULLAR",  # Turkish (significant)
        "WARUNKI I ZASADY", "ISTOTNE WARUNKI I ZASADY",  # Polish
        "BETINGELSER OG VILKÅR", "VIKTIGE BETINGELSER OG VILKÅR",  # Norwegian
        "NOTEIKUMI", "BŪTISKIE NOTEIKUMI",  # Latvian
    ]
    end = len(paragraphs)
    for i, para in enumerate(paragraphs[start + 1:], start + 1):
        para_upper = _strip_bbcode(para).upper().strip()
        if "SMS" in para_upper and "OMS" not in para_upper:
            end = i
            break
        if para_upper in tc_exact_headers:
            end = i
            break
    
    for para in paragraphs[start:end]:
        para_stripped = para.strip()
        para_upper = _strip_bbcode(para_stripped).upper()
        
        # Skip the section header itself (check against all localized markers)
        is_header = any(marker in para_upper for marker in markers)
        if is_header:
            # If header has variant, create template now
            if header_variant and not current_template:
                current_template = TemplateContent(
                    variant=header_variant,
                    send_condition=f"ClaimedReward-Template{header_variant}"
                )
                current_field = "title"
            continue
        
        # Check for explicit template variant marker
        variant_match = re.match(r"(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])", para_upper)
        if variant_match:
            if current_template:
                section.templates.append(current_template)
            variant = variant_match.group(1)
            current_template = TemplateContent(
                variant=variant,
                send_condition=f"ClaimedReward-Template{variant}"
            )
            current_field = "title"
            continue
        
        # If we hit Title/Body/CTA without a template, create default (A)
        if not current_template and (_is_standalone_label(para_upper, TITLE_LABELS) or
                                      _is_standalone_label(para_upper, BODY_LABELS) or
                                      _is_standalone_label(para_upper, CTA_LABELS) or
                                      any(para_upper.startswith(lbl) for lbl in TITLE_LABELS | BODY_LABELS | CTA_LABELS)):
            current_template = TemplateContent(
                variant="A",
                send_condition="ClaimedReward-TemplateA"
            )
        
        if current_template:
            # Handle embedded labels (e.g., "Title\nActual title text")
            para_plain_upper = _strip_bbcode(para_stripped).upper()
            if _is_label_start(para_plain_upper, TITLE_LABELS):
                current_field = "title"
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.title = rest
                continue
            elif _is_label_start(para_plain_upper, BODY_LABELS):
                current_field = "body"
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.body = rest
                continue
            elif _is_label_start(para_plain_upper, CTA_LABELS):
                current_field = "cta"
                newline_pos = para_stripped.find('\n')
                if newline_pos > 0:
                    rest = para_stripped[newline_pos + 1:].strip()
                    if rest:
                        current_template.cta = rest
                continue
            
            # Handle standalone field labels
            if _is_standalone_label(para_upper, TITLE_LABELS):
                current_field = "title"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS) is not None:
                current_field = "title"
                rest = _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS)
                if rest:
                    current_template.title = rest
                continue
            elif _is_standalone_label(para_upper, BODY_LABELS):
                current_field = "body"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, BODY_LABELS) is not None:
                current_field = "body"
                rest = _strip_label_prefix(para_stripped, para_upper, BODY_LABELS)
                if rest:
                    current_template.body = rest
                continue
            elif _is_standalone_label(para_upper, CTA_LABELS):
                current_field = "cta"
                continue
            elif _strip_label_prefix(para_stripped, para_upper, CTA_LABELS) is not None:
                current_field = "cta"
                rest = _strip_label_prefix(para_stripped, para_upper, CTA_LABELS)
                if rest:
                    current_template.cta = rest
                continue
            elif current_field and para_stripped:
                current_value = getattr(current_template, current_field) or ""
                new_value = (current_value + "\n" + para_stripped).strip()
                setattr(current_template, current_field, new_value)
    
    # Don't forget the last template
    if current_template:
        section.templates.append(current_template)
    
    return section if section.templates else None


def _parse_sms_section(paragraphs: list[str], section_type: str) -> Optional[SmsSection]:
    """Parse SMS section (Launch or Reminder) with template variants."""
    marker_key = f"{section_type.upper()}_SMS"
    markers = SECTION_MARKERS.get(marker_key, [f"{section_type.upper()} SMS"])
    
    start = _find_section_start(paragraphs, markers)

    # Fallback for table-based docs where SMS table has "SMS" header then
    # standalone "LAUNCH" or "REMINDER" sub-headers (not combined like "LAUNCH SMS").
    if start == -1:
        target = section_type.upper()  # "LAUNCH" or "REMINDER"
        # Estonian equivalents
        et_targets = {"LAUNCH": "LANSSEERIMINE", "REMINDER": "MEELDETULETUS"}
        et_target = et_targets.get(target, "")
        for i, para in enumerate(paragraphs):
            para_upper_stripped = _strip_bbcode(para).upper().strip()
            # Strip leading $ (Word doc artifact in some templates)
            if para_upper_stripped.startswith('$'):
                para_upper_stripped = para_upper_stripped[1:].strip()
            # Match "SMS", "СМС", or lines starting with "SMS " (e.g., "SMS 18+. sms.%%BrandDomain%%/mensagem")
            is_sms_header = para_upper_stripped in ("SMS", "СМС") or para_upper_stripped.startswith("SMS ") or para_upper_stripped.startswith("СМС ")
            if is_sms_header:
                # Scan ahead for the target sub-header within the SMS block
                for j in range(i + 1, len(paragraphs)):
                    line = _strip_bbcode(paragraphs[j]).upper().strip()
                    # Strip leading $ (Word doc artifact)
                    if line.startswith('$'):
                        line = line[1:].strip()
                    # Stop if we hit T&C or ADDITIONAL INFO (left the SMS table)
                    if line in ("TAC", "T&C", "T&CS", "ADDITIONAL INFO", "NOTEIKUMI", "TINGIMUSED") or line.startswith("CAMPAIGNWIZARD"):
                        break
                    if line == target or line == et_target:
                        start = j
                        break
                if start != -1:
                    break
    if start == -1:
        return None
    
    # Map section type to SendCondition
    send_condition = "JoinedCampaign" if section_type == "Launch" else "NotOptedIn"
    
    section = SmsSection(section_type=section_type)
    current_template = None
    
    # Find end of SMS section
    # Use only exact T&C section headers (not partial matches that could appear in body text)
    tc_exact_headers = [
        "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
        "ΟΡΟΙ ΚΑΙ ΠΡΟΫΠΟΘΕΣΕΙΣ", "ΣΗΜΑΝΤΙΚΟΙ ΟΡΟΙ", "ΠΛΗΡΕΙΣ ΟΡΟΙ",
        "TÉRMINOS Y CONDICIONES", "CONDICIONES IMPORTANTES",
        "TERMOS E CONDIÇÕES", "TERMOS IMPORTANTES",
        "TERMINI E CONDIZIONI",
        "CONDITIONS GÉNÉRALES",
        "ALLGEMEINE GESCHÄFTSBEDINGUNGEN", "AGB",
        "VILKÅR OG BETINGELSER", "ALLMÄNNA VILLKOR",
        "TINGIMUSED", "OLULISED TINGIMUSED", "TÄIELIKUD TINGIMUSED",  # Estonian
        "REEGLID JA TINGIMUSED", "REEGLID", "OLULISED REEGLID",  # Estonian (alt)
        "REGLUR OG SKILYRÐI", "SKILYRÐI",  # Icelandic
        "KÄYTTÖEHDOT", "EHDOT",  # Finnish
        "TERMS & CONDITIONS",  # English variant
        "ŞARTLAR VE KOŞULLAR", "ŞARTLAR & KOŞULLAR",  # Turkish
        "ÖNEMLI ŞARTLAR VE KOŞULLAR",  # Turkish (significant)
        "WARUNKI I ZASADY", "ISTOTNE WARUNKI I ZASADY",  # Polish
        "BETINGELSER OG VILKÅR", "VIKTIGE BETINGELSER OG VILKÅR",  # Norwegian
        "NOTEIKUMI", "BŪTISKIE NOTEIKUMI",  # Latvian
    ]
    end = len(paragraphs)
    for i, para in enumerate(paragraphs[start + 1:], start + 1):
        para_upper = _strip_bbcode(para).upper().strip()
        # Strip leading $ (Word doc artifact in some templates)
        if para_upper.startswith('$'):
            para_upper = para_upper[1:].strip()
        
        # Check if we hit another major section
        # Normalize multiple spaces to single for reliable substring matching
        para_upper_norm = re.sub(r'\s+', ' ', para_upper)
        if section_type == "Launch" and ("REMINDER SMS" in para_upper_norm or "REMINDER СМС" in para_upper_norm or "SMS - REMINDER" in para_upper_norm or "MEELDETULETUS SMS" in para_upper_norm or "MEELDETULETUS" in para_upper_norm or para_upper_norm in ("REMINDER", "MEELDETULETUS")):
            end = i
            break
        # Only match exact T&C section headers
        if para_upper in tc_exact_headers:
            end = i
            break
    
    for para in paragraphs[start:end]:
        para_stripped = para.strip()
        
        # Strip bullet prefix (• or - or *) if present, for matching
        # This handles Word docs where SMS templates are formatted as list items
        para_for_match = _strip_bbcode(para_stripped)
        if para_for_match.startswith("• "):
            para_for_match = para_for_match[2:]
        elif para_for_match.startswith("- ") or para_for_match.startswith("* "):
            para_for_match = para_for_match[2:]
        
        # Strip leading $ (Word doc artifact in some templates, e.g. "$REMINDER")
        if para_for_match.startswith('$'):
            para_for_match = para_for_match[1:].strip()
        
        para_upper = para_for_match.upper().strip()
        
        # Check for combined section+variant header (table-based docs)
        # e.g. "LAUNCH SMS - TemplateA", "REMINDER SMS - TemplateB"
        # Also handles Estonian: "Näidis" (Template), "Meeldetuletus" (Reminder), "Lansseerimine" (Launch)
        combined_sms_match = re.search(r"(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])", para_upper)
        has_sms = "SMS" in para_upper or "СМС" in para_upper
        if combined_sms_match and ("LAUNCH" in para_upper or "REMINDER" in para_upper or "MEELDETULETUS" in para_upper or "LANSSEERIMINE" in para_upper) and has_sms:
            if current_template:
                section.templates.append(current_template)
            # Extract any inline body after the header line (table-based docs
            # may have header + body in the same cell separated by newline)
            inline_body = None
            variant_end = combined_sms_match.end()
            rest_after_variant = para_for_match[variant_end:]
            # Check for newline - body is on the next line within the same paragraph
            newline_pos = rest_after_variant.find('\n')
            if newline_pos >= 0:
                inline_body = rest_after_variant[newline_pos + 1:].strip()
            else:
                # No newline - check if there's content directly after variant (e.g. colon+text)
                rest_stripped = rest_after_variant.strip().lstrip(':').strip()
                if rest_stripped:
                    inline_body = rest_stripped
            current_template = TemplateContent(
                variant=combined_sms_match.group(1),
                send_condition=send_condition,
                body=inline_body if inline_body else None
            )
            continue

        # Skip SMS section headers (e.g., "SMS TEMPLATES", "LAUNCH SMS", "REMINDER SMS", "SMS - LAUNCH")
        # Also handles Cyrillic "СМС" (Russian) and Estonian equivalents
        if has_sms and ("LAUNCH" in para_upper or "REMINDER" in para_upper or "MEELDETULETUS" in para_upper or "LANSSEERIMINE" in para_upper or para_upper in ("SMS", "СМС") or "TEMPLATES" in para_upper):
            continue

        # Skip standalone sub-headers (table-based docs: "LAUNCH" or "REMINDER" on their own line)
        # Also handles Estonian equivalents and $ prefixed headers
        if para_upper in ("LAUNCH", "REMINDER", "LANSSEERIMINE", "MEELDETULETUS"):
            continue

        # Skip standalone "BODY" labels (table-based docs have [Body, content] rows)
        if para_upper == "BODY" or para_upper == "BODY:":
            continue
        
        # Check for template variant marker with inline body (e.g., "Template A: body content...")
        inline_match = re.match(r"^(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])[:：\s]+(.+)$", para_upper)
        if inline_match:
            if current_template:
                section.templates.append(current_template)
            variant = inline_match.group(1)
            # Get the actual body content (not uppercased), use stripped version without bullet
            inline_body_match = re.match(r"^(?:[Tt]emplate|[Mm]all|[Pp]õhi|[Nn]äidis)\s*[A-F][:：\s]+(.+)$", para_for_match)
            if inline_body_match:
                body_content = inline_body_match.group(1)
            else:
                body_content = inline_match.group(2)
            current_template = TemplateContent(
                variant=variant,
                send_condition=send_condition,
                body=body_content
            )
            continue
        
        # Check for template variant marker without inline body (e.g., "Template A" on its own line)
        variant_match = re.match(r"^(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])$", para_upper)
        if variant_match:
            if current_template:
                section.templates.append(current_template)
            current_template = TemplateContent(
                variant=variant_match.group(1),
                send_condition=send_condition
            )
            continue
        
        # Also check for standalone letter variants
        if para_upper in TEMPLATE_VARIANTS:
            if current_template:
                section.templates.append(current_template)
            current_template = TemplateContent(
                variant=para_upper,
                send_condition=send_condition
            )
            continue
        
        if current_template and para_stripped:
            # SMS only has body content
            current_body = current_template.body or ""
            current_template.body = (current_body + " " + para_stripped).strip()
            
            # Check if this line ends the SMS (unsubscribe link pattern)
            # e.g., "sms.%%PalantirDomain%%/mc" or similar URL patterns
            if "%%PalantirDomain%%" in para_stripped and "/m" in para_stripped.lower():
                # This template is complete, save it and reset
                section.templates.append(current_template)
                current_template = None
    
    if current_template:
        section.templates.append(current_template)
    
    return section if section.templates else None


def _parse_tc_section(paragraphs: list[str]) -> Optional[TcSection]:
    """Parse Terms & Conditions section."""
    # Use strict matching to avoid matching TC markers in body text
    start = _find_tc_section_start(paragraphs, SECTION_MARKERS["TC"])
    if start == -1:
        return None
    
    section = TcSection()
    current_field = None
    
    # Localized markers for field detection
    significant_markers = [
        "SIGNIFICANT TERMS", "SIGNIFICANT T",  # English
        "SIGNIFICANT TERMS & CONDITIONS",  # Variant
        "SIGNIFICANTTERMS",  # CamelCase (table-based docs)
        "ΣΗΜΑΝΤΙΚΟΙ ΟΡΟΙ", "ΣΗΜΑΝΤΙΚΟΊ ΌΡΟΙ",  # Greek (different accent forms)
        "CONDICIONES IMPORTANTES", "TÉRMINOS IMPORTANTES",  # Spanish
        "TERMOS IMPORTANTES",  # Portuguese
        "TERMINI IMPORTANTI",  # Italian
        "CONDITIONS IMPORTANTES",  # French
        "WICHTIGE BEDINGUNGEN",  # German
        "OLULISED TINGIMUSED", "OLULISED REEGLID",  # Estonian
        "ОСНОВНЫЕ ПРАВИЛА", "ВАЖНЫЕ УСЛОВИЯ",  # Russian
        "TÄRKEÄT EHDOT",  # Finnish
        "VIKTIGE BETINGELSER OG VILKÅR", "VIKTIGE BETINGELSER",  # Norwegian
        "ISTOTNE WARUNKI I ZASADY", "ISTOTNE WARUNKI",  # Polish
        "ÖNEMLI ŞARTLAR VE KOŞULLAR", "ÖNEMLI ŞARTLAR",  # Turkish
        "VIGTIGE BETINGELSER OG VILKÅR", "VIGTIGE BETINGELSER",  # Danish
        "BŪTISKIE NOTEIKUMI",  # Latvian
    ]
    full_terms_markers = [
        "TERMS AND CONDITIONS", "TERMS & CONDITIONS", "FULL TERMS", "T&CS", "T&C",  # English
        "TERMSANDCONDITIONS",  # CamelCase (table-based docs)
        "ΠΛΗΡΕΙΣ ΟΡΟΙ", "ΠΛΉΡΕΙΣ ΌΡΟΙ", "ΟΡΟΙ ΚΑΙ ΠΡΟΫΠΟΘΕΣΕΙΣ",  # Greek
        "TÉRMINOS Y CONDICIONES", "CONDICIONES COMPLETAS",  # Spanish
        "TERMOS E CONDIÇÕES", "TERMOS COMPLETOS",  # Portuguese
        "TERMINI E CONDIZIONI",  # Italian
        "CONDITIONS GÉNÉRALES",  # French
        "ALLGEMEINE GESCHÄFTSBEDINGUNGEN", "AGB",  # German
        "TÄIELIKUD TINGIMUSED", "TINGIMUSED",  # Estonian
        "REEGLID JA TINGIMUSED", "REEGLID",  # Estonian (alt)
        "REGLUR OG SKILYRÐI",  # Icelandic
        "ПОЛНЫЕ ПРАВИЛА", "ПРАВИЛА И УСЛОВИЯ",  # Russian
        "TÄYDELLISET EHDOT", "KÄYTTÖEHDOT",  # Finnish
        "BETINGELSER OG VILKÅR",  # Norwegian
        "WARUNKI I ZASADY",  # Polish
        "ŞARTLAR VE KOŞULLAR", "ŞARTLAR & KOŞULLAR",  # Turkish
        "VILKÅR OG BETINGELSER",  # Danish
        "NOTEIKUMI",  # Latvian
    ]

    # Stop markers: sections that come after T&C
    tc_stop_markers = ["ADDITIONAL INFO", "ADDITIONAL INFORMATION"]
    
    for para in paragraphs[start:]:
        para_stripped = para.strip()
        para_upper = _strip_bbcode(para_stripped).upper()
        
        # Stop at ADDITIONAL INFO section (comes after T&C in table-based docs)
        if any(para_upper.startswith(m) for m in tc_stop_markers):
            break

        # Check for significant terms section
        if any(marker in para_upper for marker in significant_markers):
            current_field = "significant_terms"
        # Check for full terms section
        elif any(marker in para_upper for marker in full_terms_markers):
            current_field = "terms_and_conditions"
        elif current_field and para_stripped:
            current_value = getattr(section, current_field) or ""
            new_value = (current_value + "\n" + para_stripped).strip()
            setattr(section, current_field, new_value)

    # Convert bullet lines to proper BBCode lists
    if section.significant_terms:
        section.significant_terms = _bullets_to_bbcode_list(section.significant_terms)
    if section.terms_and_conditions:
        section.terms_and_conditions = _bullets_to_bbcode_list(section.terms_and_conditions)

    return section


def _parse_push_section(paragraphs: list[str], section_type: str) -> Optional[PushSection]:
    """Parse Push Notification section (Launch or Reminder) with template variants."""
    marker_key = f"{section_type.upper()}_PUSH"
    markers = SECTION_MARKERS.get(marker_key, [f"{section_type.upper()} PUSH"])

    start = _find_section_start(paragraphs, markers)
    if start == -1:
        return None

    send_condition = "JoinedCampaign" if section_type == "Launch" else "NotOptedIn"

    section = PushSection(section_type=section_type)
    current_template = None
    current_field = None

    # Find end of section
    end = len(paragraphs)
    for i, para in enumerate(paragraphs[start + 1:], start + 1):
        para_upper = _strip_bbcode(para).upper().strip()
        if section_type == "Launch" and "REMINDER" in para_upper and "PUSH" in para_upper:
            end = i
            break
        if "SMS" in para_upper and "PUSH" not in para_upper:
            end = i
            break
        if "OMS" in para_upper and "PUSH" not in para_upper:
            end = i
            break
        tc_exact_headers = [
            "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
            "TERMS & CONDITIONS",
        ]
        if para_upper in tc_exact_headers:
            end = i
            break

    _TEMPLATE_RE = re.compile(
        r"(?:PUSH(?:\s+NOTIFICATION)?)\s*[-–—]?\s*TEMPLATE\s*([A-F])",
        re.IGNORECASE,
    )
    _COMBINED_RE = re.compile(
        r"(?:LAUNCH|REMINDER)\s+PUSH(?:\s+NOTIFICATION)?\s*[-–—]\s*TEMPLATE\s*([A-F])",
        re.IGNORECASE,
    )

    for para in paragraphs[start:end]:
        para_stripped = para.strip()
        para_upper = _strip_bbcode(para_stripped).upper()

        # Check for combined section+variant header
        combined_match = _COMBINED_RE.search(para_upper)
        if combined_match:
            variant = combined_match.group(1).upper()
            current_template = TemplateContent(variant=variant, send_condition=send_condition)
            section.templates.append(current_template)
            current_field = None
            continue

        # Check for standalone template header
        template_match = _TEMPLATE_RE.search(para_upper)
        if template_match:
            variant = template_match.group(1).upper()
            current_template = TemplateContent(variant=variant, send_condition=send_condition)
            section.templates.append(current_template)
            current_field = None
            continue

        if not current_template:
            continue

        # Field labels
        if _is_standalone_label(para_upper, TITLE_LABELS):
            current_field = "title"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS)
            current_field = "title"
            if rest:
                current_template.title = rest
            continue

        if _is_standalone_label(para_upper, BODY_LABELS):
            current_field = "body"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, BODY_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, BODY_LABELS)
            current_field = "body"
            if rest:
                current_template.body = rest
            continue

        if _is_standalone_label(para_upper, ACTION_KEY_LABELS):
            current_field = "action_key"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, ACTION_KEY_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, ACTION_KEY_LABELS)
            current_field = "action_key"
            if rest:
                current_template.action_key = rest
            continue

        if _is_standalone_label(para_upper, ACTION_VALUE_LABELS):
            current_field = "action_value"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, ACTION_VALUE_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, ACTION_VALUE_LABELS)
            current_field = "action_value"
            if rest:
                current_template.action_value = rest
            continue

        # Content line — append to current field
        if current_field and para_stripped:
            current_value = getattr(current_template, current_field) or ""
            new_value = (current_value + "\n" + para_stripped).strip() if current_value else para_stripped
            setattr(current_template, current_field, new_value)

    if not section.templates:
        return None

    return section


def _parse_reward_push_section(paragraphs: list[str]) -> Optional[PushSection]:
    """Parse Reward Push Notification section with template variants.

    Each template variant gets its own send_condition: ClaimedReward-TemplateX.
    Mirrors _parse_reward_oms_section but for push notification fields.
    """
    markers = SECTION_MARKERS.get("REWARD_PUSH", ["REWARD PUSH"])

    start = _find_section_start(paragraphs, markers)
    if start == -1:
        return None

    section = PushSection(section_type="Reward")
    current_template = None
    current_field = None

    # Check if the header itself contains a variant letter
    header_variant = None
    header_text = _strip_bbcode(paragraphs[start]).upper().strip()
    variant_match = re.search(r"TEMPLATE\s*([A-F])", header_text)
    if variant_match:
        header_variant = variant_match.group(1)

    # Find end of section
    tc_exact_headers = [
        "T&C", "T&CS", "TAC", "TERMS AND CONDITIONS", "SIGNIFICANT TERMS",
        "TERMS & CONDITIONS",
    ]
    end = len(paragraphs)
    for i, para in enumerate(paragraphs[start + 1:], start + 1):
        para_upper = _strip_bbcode(para).upper().strip()
        if "SMS" in para_upper and "PUSH" not in para_upper:
            end = i
            break
        if "OMS" in para_upper and "PUSH" not in para_upper:
            end = i
            break
        if para_upper in tc_exact_headers:
            end = i
            break

    for para in paragraphs[start:end]:
        para_stripped = para.strip()
        para_upper = _strip_bbcode(para_stripped).upper()

        # Skip the section header
        is_header = any(marker in para_upper for marker in markers)
        if is_header:
            if header_variant and not current_template:
                current_template = TemplateContent(
                    variant=header_variant,
                    send_condition=f"ClaimedReward-Template{header_variant}",
                )
                current_field = "title"
            continue

        # Check for explicit template variant marker
        variant_match = re.match(r"(?:TEMPLATE|MALL|PÕHI|NÄIDIS)\s*([A-F])", para_upper)
        if variant_match:
            if current_template:
                section.templates.append(current_template)
            variant = variant_match.group(1)
            current_template = TemplateContent(
                variant=variant,
                send_condition=f"ClaimedReward-Template{variant}",
            )
            current_field = "title"
            continue

        # If we hit a field label without a template, create default (A)
        if not current_template and (
            _is_standalone_label(para_upper, TITLE_LABELS)
            or _is_standalone_label(para_upper, BODY_LABELS)
            or _is_standalone_label(para_upper, ACTION_KEY_LABELS)
            or _is_standalone_label(para_upper, ACTION_VALUE_LABELS)
            or any(para_upper.startswith(lbl) for lbl in TITLE_LABELS | BODY_LABELS | ACTION_KEY_LABELS | ACTION_VALUE_LABELS)
        ):
            current_template = TemplateContent(
                variant="A",
                send_condition="ClaimedReward-TemplateA",
            )

        if not current_template:
            continue

        # Field labels
        if _is_standalone_label(para_upper, TITLE_LABELS):
            current_field = "title"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, TITLE_LABELS)
            current_field = "title"
            if rest:
                current_template.title = rest
            continue

        if _is_standalone_label(para_upper, BODY_LABELS):
            current_field = "body"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, BODY_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, BODY_LABELS)
            current_field = "body"
            if rest:
                current_template.body = rest
            continue

        if _is_standalone_label(para_upper, ACTION_KEY_LABELS):
            current_field = "action_key"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, ACTION_KEY_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, ACTION_KEY_LABELS)
            current_field = "action_key"
            if rest:
                current_template.action_key = rest
            continue

        if _is_standalone_label(para_upper, ACTION_VALUE_LABELS):
            current_field = "action_value"
            continue
        elif _strip_label_prefix(para_stripped, para_upper, ACTION_VALUE_LABELS) is not None:
            rest = _strip_label_prefix(para_stripped, para_upper, ACTION_VALUE_LABELS)
            current_field = "action_value"
            if rest:
                current_template.action_value = rest
            continue

        # Content line — append to current field
        if current_field and para_stripped:
            current_value = getattr(current_template, current_field) or ""
            new_value = (current_value + "\n" + para_stripped).strip() if current_value else para_stripped
            setattr(current_template, current_field, new_value)

    if current_template and current_template not in section.templates:
        section.templates.append(current_template)

    if not section.templates:
        return None

    return section


def parse_documents_from_folder(folder_path: Path) -> list[ParsedDocument]:
    """
    Parse all Word documents in a folder.
    
    Args:
        folder_path: Path to folder containing .docx files
        
    Returns:
        List of ParsedDocument objects
    """
    documents = []
    for docx_file in folder_path.glob("*.docx"):
        # Skip temp files
        if docx_file.name.startswith("~"):
            continue
        try:
            parsed = parse_word_document(docx_file)
            documents.append(parsed)
        except Exception as e:
            print(f"Error parsing {docx_file.name}: {e}")
    
    return documents
